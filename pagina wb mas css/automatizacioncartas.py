import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

def set_cell_shading(cell, color):
    """
    Establece el color de fondo de una celda de tabla en Word.
    color: string hexadecimal (ej. 'FF0000')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def generar_cartas():
    # 1. Cargar los datos del Excel
    excel_file = "calificaciones_alumnos.xlsx"
    
    # Si el archivo no existe, lo creamos primero
    if not os.path.exists(excel_file):
        print(f"Generando archivo {excel_file}...")
        alumnos_list = [
            "Juan Pérez", "María García", "Luis Rodríguez", "Ana Martínez", "Carlos López",
            "Laura González", "Pedro Sánchez", "Elena Romero", "Jorge Fernández", "Sofía Ruiz"
        ]
        materias_list = ["Matemáticas", "Español", "Ciencias", "Historia", "Inglés"]
        np.random.seed(42)
        calif_data = np.random.randint(5, 11, size=(len(alumnos_list), len(materias_list)))
        df_new = pd.DataFrame(calif_data, columns=materias_list, index=alumnos_list)
        df_new.index.name = "Alumnos"
        df_new.to_excel(excel_file)
    
    # Leer el Excel
    df = pd.read_excel(excel_file, index_col=0)
    
    # Crear carpeta para las cartas si no existe
    output_dir = "cartas_alumnos"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Configuración de fecha en español
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", 
             "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    ahora = datetime.datetime.now()
    fecha_es = f"{ahora.day} de {meses[ahora.month-1]} de {ahora.year}"

    # 2. Iterar sobre cada alumno
    for nombre_alumno, row in df.iterrows():
        doc = Document()
        
        # --- Estilos ---
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # --- Fecha ---
        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fecha.add_run(f"Ciudad de México, a {fecha_es}")

        # --- Destinatario ---
        doc.add_paragraph("\nESTIMADO PADRE DE FAMILIA / TUTOR,\nPRESENTE.")

        # --- Saludo ---
        doc.add_paragraph(f"Reciba un cordial saludo. Por medio de la presente, nos dirigimos a usted para informarle sobre el desempeño académico del alumno(a) {nombre_alumno}.")

        # --- Cuerpo ---
        doc.add_paragraph("Le hacemos entrega de los resultados correspondientes al primer reporte de calificaciones del ciclo escolar actual. A continuación, se detallan las notas obtenidas en cada materia:")

        # --- Tabla de Calificaciones ---
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Materia'
        hdr_cells[1].text = 'Calificación'
        
        # Llenar tabla y aplicar colores
        suma_calif = 0
        for materia in df.columns:
            calif = row[materia]
            suma_calif += calif
            row_cells = table.add_row().cells
            row_cells[0].text = str(materia)
            row_cells[1].text = str(calif)
            
            # Sombreado: Rojo para < 6, Azul para >= 6
            color = "FF9999" if calif < 6 else "99CCFF" # Rojo claro / Azul claro
            set_cell_shading(row_cells[0], color)
            set_cell_shading(row_cells[1], color)

        promedio = suma_calif / len(df.columns)
        
        # --- Recomendación adicional si el promedio es menor a 6 ---
        p_promedio = doc.add_paragraph()
        p_promedio.add_run(f"\nPromedio General: {promedio:.1f}").bold = True
        
        if promedio < 6:
            rec = doc.add_paragraph()
            run_rec = rec.add_run("RECOMENDACIÓN: El alumno presenta un promedio insuficiente. Se recomienda establecer un horario de estudio riguroso, asistir a asesorías extraescolares y mantener una comunicación constante con los docentes para mejorar su rendimiento académico.")
            run_rec.bold = True
            run_rec.font.color.rgb = RGBColor(200, 0, 0) # Rojo

        # --- Despedida y Firma ---
        p_despedida = doc.add_paragraph("\nSin más por el momento, agradecemos su atención y quedamos a su disposición para cualquier duda o aclaración.")
        p_despedida.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p_firma = doc.add_paragraph("\n\n___________________________\nDIRECCIÓN ACADÉMICA\nINSTITUCIÓN EDUCATIVA")
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Guardar Documento ---
        safe_name = str(nombre_alumno).replace(" ", "_")
        file_path = os.path.join(output_dir, f"Carta_{safe_name}.docx")
        doc.save(file_path)
        print(f"Generada: {file_path}")

if __name__ == "__main__":
    generar_cartas()
