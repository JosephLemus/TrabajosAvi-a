import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from datetime import datetime
import os
import matplotlib.pyplot as plt
import io

def set_cell_shading(cell, color):
    """
    Establece el sombreado de una celda.
    color: string hexadecimal sin el '#' (ej: 'D9EAD3')
    """
    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def obtener_fecha_espanol():
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    ahora = datetime.now()
    dia = ahora.day
    mes = meses[ahora.month]
    anio = ahora.year
    return f"{dia} de {mes} de {anio}"

def generar_grafica(calificaciones, alumno):
    materias = list(calificaciones.keys())
    valores = list(calificaciones.values())

    plt.figure(figsize=(6, 3)) # Reducir altura para que quepa mejor
    bars = plt.bar(materias, valores, color='#4A90E2')
    plt.ylim(0, 11)
    plt.title(f'Desempeño Académico - {alumno}', fontsize=10)
    plt.ylabel('Calificación', fontsize=8)
    plt.xticks(fontsize=8)
    plt.yticks(fontsize=8)
    
    # Añadir etiquetas de valor sobre las barras
    for bar in bars:
        yval = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2, yval + 0.1, f'{yval}', ha='center', va='bottom', fontsize=8)

    # Guardar en un buffer de memoria
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=100)
    plt.close()
    buf.seek(0)
    return buf

def generar_cartas():
    # 1. Cargar los datos del Excel
    excel_path = "calificaciones_alumnos.xlsx"
    if not os.path.exists(excel_path):
        print(f"Error: El archivo {excel_path} no existe.")
        return

    df = pd.DataFrame(pd.read_excel(excel_path))
    materias = [col for col in df.columns if col != "Alumno"]

    # Carpeta de salida para los reportes
    output_folder = "reportes_word"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 2. Iterar sobre cada alumno
    for index, row in df.iterrows():
        alumno = row["Alumno"]
        calificaciones = {materia: row[materia] for materia in materias}
        promedio = sum(calificaciones.values()) / len(calificaciones)

        # Crear documento Word
        doc = Document()

        # Ajustar márgenes para que todo quepa en una hoja
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Estilo base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10) # Reducir ligeramente el tamaño de fuente

        # Fecha (Derecha)
        fecha_texto = obtener_fecha_espanol()
        p_fecha = doc.add_paragraph(fecha_texto)
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Destinatario
        p_dest = doc.add_paragraph("\nEstimado Padre de Familia / Tutor,")
        p_dest.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph("Presente.")
        
        # Saludo
        p_saludo = doc.add_paragraph("\nReciba un cordial saludo por parte de la institución académica.")
        p_saludo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Cuerpo del documento - JUSTIFICADO
        p_cuerpo = doc.add_paragraph()
        p_cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_cuerpo.add_run(f"Por medio de la presente, nos dirigimos a usted para hacerle entrega oficial de las calificaciones obtenidas por su hijo(a) ")
        p_cuerpo.add_run(f"{alumno}").bold = True
        p_cuerpo.add_run(f" correspondientes al primer reporte del ciclo escolar actual. A continuación se detalla el desempeño académico:")

        # Tabla de calificaciones con sombreado
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'Table Grid'
        
        # Cabecera con color (Azul claro)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Materia'
        hdr_cells[1].text = 'Calificación'
        for cell in hdr_cells:
            set_cell_shading(cell, "BDD7EE") # Azul claro
            cell.paragraphs[0].runs[0].bold = True

        # Filas con colores alternos o fijos
        for i, (materia, calif) in enumerate(calificaciones.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = materia
            row_cells[1].text = str(calif)
            # Sombreado alterno gris muy claro para legibilidad
            if i % 2 == 0:
                set_cell_shading(row_cells[0], "F2F2F2")
                set_cell_shading(row_cells[1], "F2F2F2")

        # Promedio
        p_promedio = doc.add_paragraph()
        p_promedio.add_run(f"\nPromedio General: {promedio:.2f}").bold = True
        p_promedio.paragraph_format.space_after = Pt(5)

        # Recomendación si promedio < 6 - JUSTIFICADO
        if promedio < 6:
            p_recomendacion = doc.add_paragraph()
            p_recomendacion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_recomendacion.add_run("RECOMENDACIÓN: ").bold = True
            p_recomendacion.add_run("Se observa que el desempeño académico actual requiere atención inmediata. Recomendamos programar una asesoría personalizada con los docentes de las materias con bajo rendimiento y establecer un plan de estudio diario en casa.")
            p_recomendacion.paragraph_format.space_after = Pt(5)

        # Despedida y firma
        p_despedida = doc.add_paragraph("Sin más por el momento, agradecemos su atención y compromiso con la educación de su hijo(a).")
        p_despedida.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_despedida.paragraph_format.space_after = Pt(10)
        
        doc.add_paragraph("Atentamente,")
        p_firma = doc.add_paragraph("\n___________________________\nDirección Académica")
        p_firma.paragraph_format.space_after = Pt(5)

        # Gráfica al Final - Tamaño ajustado para caber
        doc.add_paragraph("Visualización de desempeño:").paragraph_format.space_after = Pt(2)
        grafica_buf = generar_grafica(calificaciones, alumno)
        doc.add_picture(grafica_buf, width=Inches(3.2))
        
        # Centrar la imagen (la imagen se inserta en un párrafo)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Guardar archivo
        file_name = f"Reporte_{alumno.replace(' ', '_')}.docx"
        file_path = os.path.join(output_folder, file_name)
        doc.save(file_path)
        print(f"Generado: {file_name}")

if __name__ == "__main__":
    generar_cartas()
